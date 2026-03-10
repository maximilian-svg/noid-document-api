from pathlib import Path
from docx import Document

TEMPLATE_DIR = Path("templates")
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

def replace_tags_in_paragraph(paragraph, tag_map):
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for tag, value in tag_map.items():
        new_text = new_text.replace(f"{{{{{tag}}}}}", str(value))

    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

def render_docx(template_name: str, output_name: str, tag_map: dict[str, str]) -> str:
    template_path = TEMPLATE_DIR / template_name
    output_path = OUTPUT_DIR / output_name

    doc = Document(str(template_path))

    for p in doc.paragraphs:
        replace_tags_in_paragraph(p, tag_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_tags_in_paragraph(p, tag_map)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_tags_in_paragraph(p, tag_map)
        for p in section.footer.paragraphs:
            replace_tags_in_paragraph(p, tag_map)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_tags_in_paragraph(p, tag_map)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_tags_in_paragraph(p, tag_map)

    doc.save(str(output_path))
    return str(output_path)