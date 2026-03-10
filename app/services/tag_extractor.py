import re
from docx import Document

TAG_RE = re.compile(r"\{\{(.*?)\}\}")

def extract_tags_from_docx(path: str) -> set[str]:
    doc = Document(path)
    tags = set()

    def scan_paragraphs(paragraphs):
        for p in paragraphs:
            for match in TAG_RE.findall(p.text):
                tags.add(match)

    scan_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan_paragraphs(cell.paragraphs)

    for section in doc.sections:
        scan_paragraphs(section.header.paragraphs)
        scan_paragraphs(section.footer.paragraphs)

    return tags